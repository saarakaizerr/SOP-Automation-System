"""
Build the SOP DOCX template from scratch using python-docx.
Matches the Aged Debtor Process document structure:
  - Cover page with title + metadata table
  - Table of Contents (numbered sections, dot leaders)
  - Pre-sections (Heading 2, numbered)
  - Process Map section (Heading 1, numbered)
  - Detailed Procedure (Heading 1, numbered) with steps (Heading 3)
  - Post-sections (Heading 2, numbered)

Jinja2 / docxtpl tags are embedded as plain text in the document.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor, Cm


# ── Brand colours (matching Aged Debtor / Starboard Hotels palette) ──────────
ORANGE   = RGBColor(0xE8, 0x5C, 0x1A)   # #E85C1A
DARK     = RGBColor(0x1A, 0x1A, 0x2E)   # near-black heading
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = RGBColor(0xF8, 0xF9, 0xFA)   # table alt row
BORDER   = RGBColor(0xD1, 0xD5, 0xDB)   # table border

TEMPLATE_PATH = Path("/data/templates/sop_template.docx")
_VERSION_PATH = TEMPLATE_PATH.with_suffix(".version")
_TEMPLATE_VERSION = "17"  # increment when template structure changes

_ASSETS_DIR = Path(__file__).parent / "assets"
_HEADER_IMG = _ASSETS_DIR / "header1.jpg"
_FOOTER_IMG = _ASSETS_DIR / "footer.jpg"


def _set_run_font(run, size_pt: float, bold=False, italic=False, color=None):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color


def _set_para_spacing(para, before_pt=0, after_pt=6, line_rule=None, line_val=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(int(before_pt * 20)))
    spacing.set(qn("w:after"), str(int(after_pt * 20)))
    if line_rule and line_val:
        spacing.set(qn("w:lineRule"), line_rule)
        spacing.set(qn("w:line"), str(int(line_val * 240)))
    pPr.append(spacing)


def _para_shade(para, fill_hex: str):
    """Set paragraph background shading."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def _set_cell_bg(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _table_borders(tbl):
    """Add thin borders to all table cells."""
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "D1D5DB")
        borders.append(b)
    tblPr.append(borders)


def _add_toc_entry(doc, num: str, title_tag: str, level: int = 0):
    """
    Add a single TOC line with a dot leader tab.
    level 0 = main section  (bold, numbered, no indent)
    level 1 = sub-item      (indented, no number)
    Uses static indentation values — no Jinja2 in XML attributes.
    """
    p = doc.add_paragraph()
    p.style = "Normal"

    pPr = p._p.get_or_add_pPr()

    # Static indent based on level (no Jinja2 expressions in XML attributes)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360" if level > 0 else "0")
    ind.set(qn("w:hanging"), "0")
    pPr.append(ind)

    # Tab stop: right-aligned dot-leader at 14 cm
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), "7920")   # 14 cm ≈ 7920 twips
    tabs.append(tab)
    pPr.append(tabs)
    _set_para_spacing(p, before_pt=2, after_pt=2)

    if num:
        r_num = p.add_run(f"{num}  ")
        _set_run_font(r_num, 11, bold=True, color=DARK)

    r_title = p.add_run(title_tag)
    _set_run_font(r_title, 10 if level == 0 else 9.5, bold=(level == 0), color=DARK)

    # Dot leader tab + page ref placeholder
    r_tab = p.add_run("\t")
    _set_run_font(r_tab, 10)


def _ctrl_para(doc, tag: str):
    """Add a Jinja2 control tag paragraph (for, endfor, if, else, endif)."""
    p = doc.add_paragraph(tag)
    p.style = "Normal"
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)


def _add_page_number_to_run(run) -> None:
    """Insert a PAGE field into the given run element."""
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def _build_footer(section) -> None:
    """Add Infomate-style footer to the section (non-first pages)."""
    footer = section.footer
    footer.is_linked_to_previous = False

    # Thin horizontal separator line above the footer content (mirrors header line)
    p_sep = footer.paragraphs[0]
    p_sep.clear()
    _set_para_spacing(p_sep, before_pt=0, after_pt=0)
    pPr_sep = p_sep._p.get_or_add_pPr()
    pBdr_sep = OxmlElement("w:pBdr")
    top_b = OxmlElement("w:top")
    top_b.set(qn("w:val"), "single")
    top_b.set(qn("w:sz"), "6")
    top_b.set(qn("w:space"), "1")
    top_b.set(qn("w:color"), "AAAAAA")
    pBdr_sep.append(top_b)
    pPr_sep.append(pBdr_sep)
    p_sep.add_run(" ").font.size = Pt(1)

    # URL line + right-aligned page number
    p_url = footer.add_paragraph()
    _set_para_spacing(p_url, before_pt=2, after_pt=0)
    r_url = p_url.add_run("https://www.infomateworld.com/")
    _set_run_font(r_url, 8, color=RGBColor(0x00, 0x00, 0xCC))
    pPr_url = p_url._p.get_or_add_pPr()
    tabs_url = OxmlElement("w:tabs")
    tab_url = OxmlElement("w:tab")
    tab_url.set(qn("w:val"), "right")
    tab_url.set(qn("w:pos"), "9072")
    tabs_url.append(tab_url)
    pPr_url.append(tabs_url)
    r_tab_url = p_url.add_run("\t")
    _set_run_font(r_tab_url, 8)
    r_pgnum = p_url.add_run("")
    _set_run_font(r_pgnum, 8, color=DARK)
    _add_page_number_to_run(r_pgnum)

    # Address line + right-aligned Infomate logo image
    p_addr = footer.add_paragraph()
    _set_para_spacing(p_addr, before_pt=0, after_pt=0)
    r_addr = p_addr.add_run("No 04, Leyden Bastian Street, Colombo 01, Sri Lanka")
    _set_run_font(r_addr, 8, color=DARK)
    pPr_addr = p_addr._p.get_or_add_pPr()
    tabs_addr = OxmlElement("w:tabs")
    tab_addr = OxmlElement("w:tab")
    tab_addr.set(qn("w:val"), "right")
    tab_addr.set(qn("w:pos"), "9072")
    tabs_addr.append(tab_addr)
    pPr_addr.append(tabs_addr)
    r_tab_addr = p_addr.add_run("\t")
    _set_run_font(r_tab_addr, 8)
    if _FOOTER_IMG.exists():
        r_logo = p_addr.add_run()
        r_logo.add_picture(str(_FOOTER_IMG), height=Inches(0.22))
    else:
        r_icon = p_addr.add_run("■")
        _set_run_font(r_icon, 9, bold=True, color=ORANGE)
        r_brand = p_addr.add_run("infomate")
        _set_run_font(r_brand, 8, bold=True, color=DARK)


def _build_header(section) -> None:
    """Add Infomate-style header to the section (non-first pages)."""
    header = section.header
    header.is_linked_to_previous = False

    p_hdr = header.paragraphs[0]
    p_hdr.clear()
    _set_para_spacing(p_hdr, before_pt=0, after_pt=0)
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Explicitly clear any paragraph shading inherited from the Header style
    pPr_hdr = p_hdr._p.get_or_add_pPr()
    existing_shd = pPr_hdr.find(qn("w:shd"))
    if existing_shd is not None:
        pPr_hdr.remove(existing_shd)
    clear_shd = OxmlElement("w:shd")
    clear_shd.set(qn("w:val"), "clear")
    clear_shd.set(qn("w:color"), "auto")
    clear_shd.set(qn("w:fill"), "auto")
    pPr_hdr.append(clear_shd)

    if _HEADER_IMG.exists():
        run = p_hdr.add_run()
        run.add_picture(str(_HEADER_IMG), height=Inches(0.45))
    else:
        r_hdr = p_hdr.add_run("  ")
        _set_run_font(r_hdr, 4)

    # Thin horizontal separator line below the header image
    p_line = header.add_paragraph()
    _set_para_spacing(p_line, before_pt=0, after_pt=0)
    pPr_line = p_line._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom_b = OxmlElement("w:bottom")
    bottom_b.set(qn("w:val"), "single")
    bottom_b.set(qn("w:sz"), "6")
    bottom_b.set(qn("w:space"), "1")
    bottom_b.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom_b)
    pPr_line.append(pBdr)
    r_line = p_line.add_run(" ")
    _set_run_font(r_line, 1)


def build(force: bool = False):
    if not force and TEMPLATE_PATH.exists():
        # Skip rebuild only if version matches
        current = _VERSION_PATH.read_text().strip() if _VERSION_PATH.exists() else ""
        if current == _TEMPLATE_VERSION:
            return

    TEMPLATE_PATH.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # ── Section 1: Cover page (zero margins, full-bleed) ─────────────────────
    section1 = doc.sections[0]
    section1.page_width    = Cm(21)
    section1.page_height   = Cm(29.7)
    section1.top_margin    = Cm(0)
    section1.bottom_margin = Cm(0)
    section1.left_margin   = Cm(0)
    section1.right_margin  = Cm(0)

    # ── Default paragraph style ───────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # ── COVER PAGE ────────────────────────────────────────────────────────────
    # Cover image is generated by Pillow in doc_renderer and injected via Jinja2
    p_cover = doc.add_paragraph("{{ cover_page }}")
    p_cover.style = "Normal"
    p_cover.paragraph_format.space_before = Pt(0)
    p_cover.paragraph_format.space_after = Pt(0)

    # ── Section 2: Content pages (normal margins, header/footer) ─────────────
    # NEW_PAGE section break replaces the old doc.add_page_break() after cover
    section2 = doc.add_section(WD_SECTION.NEW_PAGE)
    section2.page_width    = Cm(21)
    section2.page_height   = Cm(29.7)
    section2.top_margin    = Cm(2.5)
    section2.bottom_margin = Cm(2.5)
    section2.left_margin   = Cm(3.0)
    section2.right_margin  = Cm(2.0)
    _build_header(section2)
    _build_footer(section2)

    # ── TABLE OF CONTENTS ─────────────────────────────────────────────────────
    p_toc_head = doc.add_heading("Table of Contents", level=1)
    _set_para_spacing(p_toc_head, before_pt=0, after_pt=12)
    for run in p_toc_head.runs:
        _set_run_font(run, 16, bold=True, color=ORANGE)

    # Jinja2 loop — use if/else to select main vs sub-item paragraph style.
    # This avoids putting Jinja2 expressions inside XML attributes (unreliable).
    _ctrl_para(doc, "{%- for entry in toc_entries %}")
    _ctrl_para(doc, "{%- if not entry.is_sub %}")
    _add_toc_entry(doc, num="{{ entry.num }}", title_tag="{{ entry.title }}", level=0)
    _ctrl_para(doc, "{%- else %}")
    _add_toc_entry(doc, num="{{ entry.num }}", title_tag="{{ entry.title }}", level=1)
    _ctrl_para(doc, "{%- endif %}")
    _ctrl_para(doc, "{%- endfor %}")

    doc.add_page_break()

    # ── PROCEDURE DESCRIPTION (section 1) with sub-sections ─────────────────
    h1_proc = doc.add_heading("1  Procedure Description", level=1)
    for run in h1_proc.runs:
        _set_run_font(run, 15, bold=True, color=ORANGE)
    _set_para_spacing(h1_proc, before_pt=12, after_pt=8)

    _ctrl_para(doc, "{%- for section in procedure_sub_sections %}")
    h2_sub = doc.add_heading("{{ section.section_title }}", level=2)
    for run in h2_sub.runs:
        _set_run_font(run, 12, bold=True, color=DARK)
    _set_para_spacing(h2_sub, before_pt=10, after_pt=4)
    p_proc_sub = doc.add_paragraph("{{r section.content_text }}")
    p_proc_sub.style = "Normal"
    _ctrl_para(doc, "{%- endfor %}")

    # ── OTHER PRE-SECTIONS (Training Prerequisites, Software Applications, etc.)
    _ctrl_para(doc, "{%- for section in other_pre_sections %}")
    h1_pre = doc.add_heading("{{ section.num }}  {{ section.section_title }}", level=1)
    for run in h1_pre.runs:
        _set_run_font(run, 15, bold=True, color=ORANGE)
    _set_para_spacing(h1_pre, before_pt=12, after_pt=8)
    p_pre_content = doc.add_paragraph("{{r section.content_text }}")
    p_pre_content.style = "Normal"
    _ctrl_para(doc, "{%- endfor %}")

    doc.add_page_break()

    # ── PROCESS MAP ───────────────────────────────────────────────────────────
    h1_pm = doc.add_heading("{{ pm_section_num }}  Process Map", level=1)
    for run in h1_pm.runs:
        _set_run_font(run, 15, bold=True, color=ORANGE)
    _set_para_spacing(h1_pm, before_pt=12, after_pt=8)

    p_pm = doc.add_paragraph("{%- if process_map %}{{ process_map }}{%- else %}(Process map not configured — use the Process Map tab to build one){%- endif %}")
    p_pm.style = "Normal"

    doc.add_page_break()

    # ── DETAILED PROCEDURE ────────────────────────────────────────────────────
    h1_dp = doc.add_heading("{{ dp_section_num }}  Detailed Procedure", level=1)
    for run in h1_dp.runs:
        _set_run_font(run, 15, bold=True, color=ORANGE)
    _set_para_spacing(h1_dp, before_pt=12, after_pt=8)

    _ctrl_para(doc, "{%- for step in steps %}")

    # Step heading
    h3 = doc.add_heading("Step {{ step.sequence }}: {{ step.title }}", level=3)
    for run in h3.runs:
        _set_run_font(run, 12, bold=True, color=DARK)
    _set_para_spacing(h3, before_pt=16, after_pt=4)

    # Description
    p_desc = doc.add_paragraph("{{ step.description | default('') }}")
    p_desc.style = "Normal"
    _set_para_spacing(p_desc, before_pt=0, after_pt=4)

    # Sub-steps
    _ctrl_para(doc, "{%- for sub in step.sub_steps %}")
    p_sub = doc.add_paragraph("{{ sub }}")
    p_sub.style = "List Bullet"
    _set_para_spacing(p_sub, before_pt=0, after_pt=2)
    _ctrl_para(doc, "{%- endfor %}")

    # Screenshot — image first, bold centered caption directly below
    _ctrl_para(doc, "{%- if step.screenshot %}")
    p_ss = doc.add_paragraph("{{ step.screenshot }}")
    p_ss.style = "Normal"
    _set_para_spacing(p_ss, before_pt=6, after_pt=2)
    p_ss_label = doc.add_paragraph("Screenshot {{ step.sequence }}")
    p_ss_label.style = "Normal"
    p_ss_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p_ss_label.runs:
        _set_run_font(run, 11, bold=True, color=DARK)
    _set_para_spacing(p_ss_label, before_pt=0, after_pt=10)
    _ctrl_para(doc, "{%- endif %}")

    # Callouts
    _ctrl_para(doc, "{%- if step.callouts %}")

    p_callout_head = doc.add_paragraph("Callout References")
    p_callout_head.style = "Normal"
    for run in p_callout_head.runs:
        _set_run_font(run, 10, bold=True, italic=True, color=DARK)
    _set_para_spacing(p_callout_head, before_pt=4, after_pt=2)

    _ctrl_para(doc, "{%- for callout in step.callouts %}")
    p_cl = doc.add_paragraph("{{ callout.callout_number }}. {{ callout.label }}")
    p_cl.style = "List Number"
    _set_para_spacing(p_cl, before_pt=0, after_pt=2)
    _ctrl_para(doc, "{%- endfor %}")
    _ctrl_para(doc, "{%- endif %}")

    # Step separator
    p_sep = doc.add_paragraph()
    p_sep.style = "Normal"
    _para_shade(p_sep, "F3F4F6")
    _set_para_spacing(p_sep, before_pt=12, after_pt=0)
    r = p_sep.add_run(" ")
    _set_run_font(r, 3)

    _ctrl_para(doc, "{%- endfor %}")

    doc.add_page_break()

    # ── POST-SECTIONS ─────────────────────────────────────────────────────────
    _ctrl_para(doc, "{%- for section in sections_post %}")

    h2_post = doc.add_heading("{{ section.num }}  {{ section.section_title }}", level=2)
    for run in h2_post.runs:
        _set_run_font(run, 13, bold=True, color=ORANGE)
    _set_para_spacing(h2_post, before_pt=18, after_pt=6)

    p_post = doc.add_paragraph("{{r section.content_text }}")
    p_post.style = "Normal"

    _ctrl_para(doc, "{%- endfor %}")

    # ── SOP AUTHOR/REVIEWER/APPROVER CERTIFICATION ───────────────────────────
    h1_cert = doc.add_heading(
        "{{ cert_section_num }}  SOP Author/Reviewer/Approver Certification", level=1
    )
    for run in h1_cert.runs:
        _set_run_font(run, 15, bold=True, color=ORANGE)
    _set_para_spacing(h1_cert, before_pt=18, after_pt=8)

    cert_tbl = doc.add_table(rows=4, cols=3)
    _table_borders(cert_tbl)
    cert_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    cert_headers = ["Role", "Name", "Signature & Date"]
    cert_roles   = ["Author / SOP Writer", "Reviewer", "Approver"]

    for i, h in enumerate(cert_headers):
        cell = cert_tbl.rows[0].cells[i]
        _set_cell_bg(cell, "E85C1A")
        run = cell.paragraphs[0].add_run(h)
        _set_run_font(run, 10, bold=True, color=WHITE)

    for r_idx, role in enumerate(cert_roles):
        bg = "F8F9FA" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate([role, "", ""]):
            cell = cert_tbl.rows[r_idx + 1].cells[c_idx]
            _set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(val)
            _set_run_font(run, 10)

    doc.save(str(TEMPLATE_PATH))
    _VERSION_PATH.write_text(_TEMPLATE_VERSION)
    print(f"Template v{_TEMPLATE_VERSION} written to {TEMPLATE_PATH}")


if __name__ == "__main__":
    build(force=True)
