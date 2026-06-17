"""
Webinar Summary DOCX template — CloudNavision branded.
Indigo/purple accent scheme. Professional structure with:
  - Branded header (CN logo + title)
  - Session metadata table
  - Session overview
  - Topics covered (numbered, content, key points, screenshot)
  - Key Takeaways summary table
  - Resources & References
  - Post-Webinar Action Items
  - Sign-off footer
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ── Brand palette ─────────────────────────────────────────────────────────────
INDIGO      = RGBColor(0x43, 0x38, 0xCA)   # #4338CA — webinar accent
INDIGO_HEX  = "4338CA"
VIOLET      = RGBColor(0x6D, 0x28, 0xD9)   # #6D28D9 — deep purple
VIOLET_HEX  = "6D28D9"
INDIGO_LT   = RGBColor(0xED, 0xE9, 0xFE)   # #EDE9FE — lavender bg
INDIGO_LT_HEX  = "EDE9FE"
INDIGO_MED_HEX = "C4B5FD"                   # medium lavender for alt rows
NAVY        = RGBColor(0x0A, 0x16, 0x28)   # #0A1628 — near-black
NAVY_HEX    = "0A1628"
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
WHITE_HEX   = "FFFFFF"
DARK        = RGBColor(0x1E, 0x29, 0x3B)   # #1E293B — body text
DARK_HEX    = "1E293B"
SLATE       = RGBColor(0x64, 0x74, 0x8B)   # #64748B — muted
TEAL        = RGBColor(0x06, 0xB6, 0xD4)   # #06B6D4 — CN primary (footer)
TEAL_HEX    = "06B6D4"
GREEN       = RGBColor(0x16, 0xA3, 0x4A)   # completed status
ORANGE      = RGBColor(0xF9, 0x73, 0x16)   # pending status

TEMPLATE_PATH     = Path("/data/templates/webinar_template.docx")
_VERSION_PATH     = TEMPLATE_PATH.with_suffix(".version")
_TEMPLATE_VERSION = "5"

_ASSETS_DIR = Path(__file__).parent / "assets"
_CN_LOGO    = _ASSETS_DIR / "cloudnavision_logo.jpg"


# ── Low-level helpers ─────────────────────────────────────────────────────────

def _font(run, size: float, bold=False, italic=False, color=None):
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.name   = "Calibri"
    if color:
        run.font.color.rgb = color


def _spacing(para, before=0, after=6):
    pPr = para._p.get_or_add_pPr()
    sp  = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(int(before * 20)))
    sp.set(qn("w:after"),  str(int(after  * 20)))
    pPr.append(sp)


def _shade_para(para, fill_hex: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    pPr.append(shd)


def _cell_bg(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def _tbl_borders(tbl, color_hex="C4B5FD"):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color_hex)
        borders.append(b)
    tblPr.append(borders)


def _border_bottom(para, color_hex=INDIGO_HEX, size=8):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(size))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)


def _ctrl(doc, tag: str):
    p = doc.add_paragraph(tag)
    p.style = "Normal"
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)


def _section_bar(doc, label: str, subtitle: str = ""):
    """Full-width navy bar + indigo accent line."""
    p = doc.add_paragraph()
    _shade_para(p, NAVY_HEX)
    _spacing(p, before=14, after=0)
    r = p.add_run(f"  {label}")
    _font(r, 10, bold=True, color=WHITE)
    if subtitle:
        r2 = p.add_run(f"  —  {subtitle}")
        _font(r2, 9, italic=True, color=INDIGO_LT)
    p2 = doc.add_paragraph()
    _shade_para(p2, INDIGO_HEX)
    _spacing(p2, before=0, after=6)
    p2.add_run("").font.size = Pt(2)


def _add_page_number(run):
    for tag, text in [("begin", None), (None, " PAGE "), ("end", None)]:
        if tag:
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), tag)
            run._r.append(fc)
        else:
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = text
            run._r.append(instr)


def _build_header(section):
    """Branded header: CN logo + title + indigo separator line."""
    hdr = section.header
    hdr.is_linked_to_previous = False

    p = hdr.paragraphs[0]
    p.clear()
    _spacing(p, before=0, after=0)

    if _CN_LOGO.exists():
        r_logo = p.add_run()
        r_logo.add_picture(str(_CN_LOGO), height=Cm(1.2))
        r_space = p.add_run("  ")
        _font(r_space, 8)

    r_title = p.add_run("Webinar Summary")
    _font(r_title, 9, italic=True, color=SLATE)

    p2 = hdr.add_paragraph()
    _spacing(p2, before=0, after=0)
    pPr = p2._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), INDIGO_HEX)
    pBdr.append(bot)
    pPr.append(pBdr)
    p2.add_run(" ").font.size = Pt(1)


def _build_footer(section):
    """Footer: company name left, page number right, teal separator."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.clear()
    _spacing(p, before=4, after=0)

    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top  = OxmlElement("w:top")
    top.set(qn("w:val"),   "single")
    top.set(qn("w:sz"),    "6")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), INDIGO_HEX)
    pBdr.append(top)
    pPr.append(pBdr)

    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), "8640")
    tabs.append(tab)
    pPr.append(tabs)

    r_left = p.add_run("CloudNavision Private Limited  |  cloudnavision.com")
    _font(r_left, 8, italic=True, color=SLATE)
    r_tab = p.add_run("\t")
    _font(r_tab, 8)
    r_pg = p.add_run("Page ")
    _font(r_pg, 8, color=SLATE)
    r_num = p.add_run("")
    _font(r_num, 8, color=INDIGO)
    _add_page_number(r_num)


# ── Builder ───────────────────────────────────────────────────────────────────

def build(force: bool = False):
    if not force and TEMPLATE_PATH.exists():
        try:
            if _VERSION_PATH.read_text().strip() == _TEMPLATE_VERSION:
                return
        except Exception:
            pass

    TEMPLATE_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.top_margin    = Cm(2.8)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)
    _build_header(sec)
    _build_footer(sec)

    # ── COVER TITLE BLOCK ─────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    _shade_para(p_title, NAVY_HEX)
    _spacing(p_title, before=0, after=0)
    r = p_title.add_run("  WEBINAR SUMMARY")
    _font(r, 24, bold=True, color=WHITE)

    p_accent = doc.add_paragraph()
    _shade_para(p_accent, INDIGO_HEX)
    _spacing(p_accent, before=0, after=0)
    p_accent.add_run("").font.size = Pt(3)

    p_sub = doc.add_paragraph()
    _shade_para(p_sub, INDIGO_LT_HEX)
    _spacing(p_sub, before=0, after=0)
    r = p_sub.add_run("  {{ doc_title }}")
    _font(r, 13, bold=True, color=NAVY)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ── SESSION DETAILS TABLE ─────────────────────────────────────────────────
    _section_bar(doc, "SESSION DETAILS")

    det = doc.add_table(rows=5, cols=4)
    det.alignment = WD_TABLE_ALIGNMENT.LEFT
    _tbl_borders(det)

    col_widths = [Cm(4.0), Cm(6.5), Cm(3.5), Cm(5.0)]
    for row in det.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]

    details_data = [
        ("Session Title",   "{{ doc_title }}",     "Reference No.",  "WS-{{ generated_date | replace(' ', '') }}"),
        ("Date",            "{{ session_date }}",  "Duration",       ""),
        ("Platform / Host", "",                    "Recording",      "Available"),
        ("Presenter",       "{{ presenter }}",     "Prepared By",    "CloudNavision"),
        ("Target Audience", "",                    "Status",         "Completed"),
    ]
    for i, (l1, v1, l2, v2) in enumerate(details_data):
        row = det.rows[i]
        bg  = INDIGO_LT_HEX if i % 2 == 0 else WHITE_HEX
        _cell_bg(row.cells[0], INDIGO_MED_HEX)
        _cell_bg(row.cells[1], bg)
        _cell_bg(row.cells[2], INDIGO_MED_HEX)
        _cell_bg(row.cells[3], bg)
        r1 = row.cells[0].paragraphs[0].add_run(l1)
        _font(r1, 9, bold=True, color=NAVY)
        r2 = row.cells[1].paragraphs[0].add_run(v1)
        _font(r2, 9, color=DARK)
        r3 = row.cells[2].paragraphs[0].add_run(l2)
        _font(r3, 9, bold=True, color=NAVY)
        r4 = row.cells[3].paragraphs[0].add_run(v2)
        _font(r4, 9, color=DARK)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── SESSION OVERVIEW ──────────────────────────────────────────────────────
    _ctrl(doc, "{%- if overview_text %}")
    _section_bar(doc, "SESSION OVERVIEW")
    p_ov = doc.add_paragraph()
    r = p_ov.add_run("{{ overview_text }}")
    _font(r, 10, color=DARK)
    _spacing(p_ov, before=4, after=8)
    _ctrl(doc, "{%- endif %}")

    # ── LEARNING OBJECTIVES (static placeholder) ──────────────────────────────
    _section_bar(doc, "LEARNING OBJECTIVES")

    obj_tbl = doc.add_table(rows=3, cols=2)
    obj_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _tbl_borders(obj_tbl)
    obj_widths = [Cm(1.5), Cm(16.5)]
    objectives = [
        "Understand the end-to-end process covered in this session",
        "Identify key steps, decisions, and system interactions",
        "Apply the demonstrated workflow in day-to-day operations",
    ]
    for i, obj in enumerate(objectives):
        row = obj_tbl.rows[i]
        bg  = INDIGO_LT_HEX if i % 2 == 0 else WHITE_HEX
        row.cells[0].width = obj_widths[0]
        row.cells[1].width = obj_widths[1]
        _cell_bg(row.cells[0], INDIGO_HEX)
        _cell_bg(row.cells[1], bg)
        r_num = row.cells[0].paragraphs[0].add_run(f"  {i + 1}")
        _font(r_num, 10, bold=True, color=WHITE)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_obj = row.cells[1].paragraphs[0].add_run(obj)
        _font(r_obj, 10, color=DARK)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── TOPICS COVERED ────────────────────────────────────────────────────────
    _section_bar(doc, "TOPICS COVERED")

    _ctrl(doc, "{%- for topic in topics %}")

    # Topic heading with indigo left badge + title
    p_topic = doc.add_paragraph()
    _shade_para(p_topic, INDIGO_LT_HEX)
    _spacing(p_topic, before=10, after=0)
    r_num   = p_topic.add_run("  {{ topic.number }}.")
    _font(r_num, 11, bold=True, color=INDIGO)
    r_title = p_topic.add_run("  {{ topic.title }}")
    _font(r_title, 11, bold=True, color=NAVY)
    _border_bottom(p_topic, INDIGO_HEX, size=6)

    # Content / Description
    p_lbl = doc.add_paragraph()
    _spacing(p_lbl, before=6, after=0)
    r = p_lbl.add_run("Topic Summary:")
    _font(r, 9, bold=True, color=INDIGO)

    p_content = doc.add_paragraph()
    r = p_content.add_run("{{ topic.content }}")
    _font(r, 10, color=DARK)
    _spacing(p_content, before=2, after=4)

    # Key Points
    _ctrl(doc, "{%- if topic.key_points %}")
    p_kl = doc.add_paragraph()
    _spacing(p_kl, before=4, after=0)
    r = p_kl.add_run("  Key Points:")
    _font(r, 9, bold=True, color=WHITE)
    _shade_para(p_kl, INDIGO_HEX)

    _ctrl(doc, "{%- for point in topic.key_points %}")
    p_kp = doc.add_paragraph()
    _spacing(p_kp, before=1, after=1)
    r_bullet = p_kp.add_run("  ▸  ")
    _font(r_bullet, 10, color=INDIGO)
    r_point = p_kp.add_run("{{ point }}")
    _font(r_point, 10, color=DARK)
    _ctrl(doc, "{%- endfor %}")
    _ctrl(doc, "{%- endif %}")

    # Screenshot
    _ctrl(doc, "{%- if topic.screenshot %}")
    p_ss = doc.add_paragraph()
    r = p_ss.add_run("{{ topic.screenshot }}")
    _font(r, 10)
    _spacing(p_ss, before=6, after=2)
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_cap.add_run("Screenshot — Topic {{ topic.number }}: {{ topic.title }}")
    _font(r, 9, italic=True, color=SLATE)
    _spacing(p_cap, before=0, after=6)
    _ctrl(doc, "{%- endif %}")

    # Topic separator
    p_sep = doc.add_paragraph()
    _spacing(p_sep, before=6, after=0)
    p_sep.add_run("").font.size = Pt(1)

    _ctrl(doc, "{%- endfor %}")

    # ── KEY TAKEAWAYS SUMMARY ─────────────────────────────────────────────────
    doc.add_page_break()
    _section_bar(doc, "KEY TAKEAWAYS SUMMARY")

    _ctrl(doc, "{%- for topic in topics %}")

    # Topic heading row
    p_kt_hd = doc.add_paragraph()
    _shade_para(p_kt_hd, INDIGO_LT_HEX)
    _spacing(p_kt_hd, before=8, after=0)
    r_kt_n = p_kt_hd.add_run("  {{ topic.number }}.")
    _font(r_kt_n, 10, bold=True, color=INDIGO)
    r_kt_t = p_kt_hd.add_run("  {{ topic.title }}")
    _font(r_kt_t, 10, bold=True, color=NAVY)
    _border_bottom(p_kt_hd, INDIGO_HEX, size=4)

    _ctrl(doc, "{%- if topic.key_points %}")
    _ctrl(doc, "{%- for point in topic.key_points %}")

    p_kt = doc.add_paragraph()
    _spacing(p_kt, before=2, after=2)
    r_kt_bul = p_kt.add_run("    ▸  ")
    _font(r_kt_bul, 10, color=INDIGO)
    r_kt_pt = p_kt.add_run("{{ point }}")
    _font(r_kt_pt, 10, color=DARK)

    _ctrl(doc, "{%- endfor %}")
    _ctrl(doc, "{%- endif %}")

    _ctrl(doc, "{%- endfor %}")

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ── RESOURCES & REFERENCES ────────────────────────────────────────────────
    _ctrl(doc, "{%- if resource_sections %}")
    _section_bar(doc, "RESOURCES & REFERENCES")
    _ctrl(doc, "{%- for section in resource_sections %}")
    p_st = doc.add_paragraph()
    r = p_st.add_run("{{ section.section_title }}")
    _font(r, 10, bold=True, color=NAVY)
    _spacing(p_st, before=8, after=2)
    p_sc = doc.add_paragraph()
    r = p_sc.add_run("{{ section.content_text }}")
    _font(r, 10, color=DARK)
    _ctrl(doc, "{%- endfor %}")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    _ctrl(doc, "{%- endif %}")

    # ── POST-WEBINAR ACTION ITEMS ─────────────────────────────────────────────
    _section_bar(doc, "POST-WEBINAR ACTION ITEMS")

    pa_tbl = doc.add_table(rows=4, cols=4)
    pa_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _tbl_borders(pa_tbl)

    pa_col_widths = [Cm(1.5), Cm(8.5), Cm(4.0), Cm(4.0)]
    pa_headers    = ["#", "Action Item", "Owner", "Target Date"]
    for i, hdr_txt in enumerate(pa_headers):
        cell = pa_tbl.rows[0].cells[i]
        cell.width = pa_col_widths[i]
        _cell_bg(cell, NAVY_HEX)
        r = cell.paragraphs[0].add_run(hdr_txt)
        _font(r, 9, bold=True, color=WHITE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i in range(1, 4):
        row = pa_tbl.rows[i]
        bg  = INDIGO_LT_HEX if i % 2 != 0 else WHITE_HEX
        for j, w in enumerate(pa_col_widths):
            row.cells[j].width = w
            _cell_bg(row.cells[j], bg)
        r_num = row.cells[0].paragraphs[0].add_run(str(i))
        _font(r_num, 9, bold=True, color=INDIGO)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ── FEEDBACK & EVALUATION ─────────────────────────────────────────────────
    _section_bar(doc, "FEEDBACK & EVALUATION")

    fb_tbl = doc.add_table(rows=5, cols=3)
    fb_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _tbl_borders(fb_tbl)

    fb_col_widths = [Cm(7.0), Cm(5.0), Cm(6.0)]
    fb_headers    = ["Evaluation Criteria", "Rating (1–5)", "Comments"]
    for i, hdr_txt in enumerate(fb_headers):
        cell = fb_tbl.rows[0].cells[i]
        cell.width = fb_col_widths[i]
        _cell_bg(cell, NAVY_HEX)
        r = cell.paragraphs[0].add_run(hdr_txt)
        _font(r, 9, bold=True, color=WHITE)

    fb_criteria = [
        "Content Relevance & Clarity",
        "Presenter's Delivery",
        "Visual Aids & Screenshots",
        "Overall Session Value",
    ]
    for i, crit in enumerate(fb_criteria):
        row = fb_tbl.rows[i + 1]
        bg  = INDIGO_LT_HEX if i % 2 == 0 else WHITE_HEX
        for j, w in enumerate(fb_col_widths):
            row.cells[j].width = w
            _cell_bg(row.cells[j], bg)
        _cell_bg(row.cells[0], INDIGO_MED_HEX)
        r = row.cells[0].paragraphs[0].add_run(crit)
        _font(r, 9, bold=True, color=NAVY)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── GENERATED NOTE ────────────────────────────────────────────────────────
    p_gen = doc.add_paragraph()
    p_gen.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p_gen.add_run("Generated by CloudNavision SOP Platform  |  {{ generated_date }}")
    _font(r, 8, italic=True, color=SLATE)

    doc.save(str(TEMPLATE_PATH))
    _VERSION_PATH.write_text(_TEMPLATE_VERSION)
    print(f"Webinar template v{_TEMPLATE_VERSION} written to {TEMPLATE_PATH}")


if __name__ == "__main__":
    build(force=True)
