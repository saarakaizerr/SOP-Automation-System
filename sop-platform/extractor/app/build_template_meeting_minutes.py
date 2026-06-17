"""
Meeting Minutes DOCX template — CloudNavision branded.
Teal colour scheme. Professional structure with:
  - Branded header (CN logo + title)
  - Meeting metadata table
  - Attendees table
  - Agenda & Discussion with per-item action items
  - Consolidated Action Items Summary
  - Next Meeting section
  - Sign-off footer
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Inches

# ── Brand palette ─────────────────────────────────────────────────────────────
TEAL       = RGBColor(0x06, 0xB6, 0xD4)   # #06B6D4 — CloudNavision primary
TEAL_HEX   = "06B6D4"
NAVY       = RGBColor(0x0A, 0x16, 0x28)   # #0A1628 — near-black
NAVY_HEX   = "0A1628"
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
WHITE_HEX  = "FFFFFF"
TEAL_LT    = RGBColor(0xE0, 0xF7, 0xFA)   # #E0F7FA — light teal for alt rows
TEAL_LT_HEX = "E0F7FA"
TEAL_MED_HEX = "B2EBF2"                   # slightly darker alt row
SLATE      = RGBColor(0x64, 0x74, 0x8B)   # #64748B — muted grey
DARK       = RGBColor(0x1E, 0x29, 0x3B)   # #1E293B — body text
DARK_HEX   = "1E293B"
ORANGE     = RGBColor(0xF9, 0x73, 0x16)   # #F97316 — open/pending status
GREEN      = RGBColor(0x16, 0xA3, 0x4A)   # #16A34A — closed status

TEMPLATE_PATH     = Path("/data/templates/meeting_minutes_template.docx")
_VERSION_PATH     = TEMPLATE_PATH.with_suffix(".version")
_TEMPLATE_VERSION = "5"

_ASSETS_DIR = Path(__file__).parent / "assets"
_CN_LOGO    = _ASSETS_DIR / "cloudnavision_logo.jpg"


# ── Low-level helpers ─────────────────────────────────────────────────────────

def _font(run, size: float, bold=False, italic=False, color=None):
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.name  = "Calibri"
    if color:
        run.font.color.rgb = color


def _spacing(para, before=0, after=6):
    pPr = para._p.get_or_add_pPr()
    sp  = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(int(before * 20)))
    sp.set(qn("w:after"),  str(int(after  * 20)))
    pPr.append(sp)


def _shade_para(para, fill_hex: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    pPr.append(shd)


def _cell_bg(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def _cell_valign(cell, align="center"):
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)


def _tbl_borders(tbl, color_hex="B2EBF2"):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color_hex)
        borders.append(b)
    tblPr.append(borders)


def _border_bottom(para, color_hex=TEAL_HEX, size=8):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(size))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)


def _ctrl(doc, tag: str):
    p = doc.add_paragraph(tag)
    p.style = "Normal"
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)


def _section_bar(doc, label: str, subtitle: str = ""):
    """Full-width navy section header bar."""
    p = doc.add_paragraph()
    _shade_para(p, NAVY_HEX)
    _spacing(p, before=14, after=0)
    r = p.add_run(f"  {label}")
    _font(r, 10, bold=True, color=WHITE)
    if subtitle:
        r2 = p.add_run(f"  —  {subtitle}")
        _font(r2, 9, italic=True, color=TEAL_LT)
    # teal bottom accent line
    p2 = doc.add_paragraph()
    _shade_para(p2, TEAL_HEX)
    _spacing(p2, before=0, after=6)
    p2.add_run("").font.size = Pt(2)


def _add_page_number(run):
    for tag, text in [("begin", None), (None, " PAGE "), ("end", None)]:
        if tag:
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), tag)
            run._r.append(fc)
        else:
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = text
            run._r.append(instr)


def _build_header(section):
    """Branded header: CN logo right, teal bottom line."""
    hdr = section.header
    hdr.is_linked_to_previous = False

    p = hdr.paragraphs[0]
    p.clear()
    _spacing(p, before=0, after=0)

    if _CN_LOGO.exists():
        r_logo = p.add_run()
        r_logo.add_picture(str(_CN_LOGO), height=Cm(1.2))
        r_space = p.add_run("  ")
        _font(r_space, 8)

    r_title = p.add_run("Meeting Minutes")
    _font(r_title, 9, italic=True, color=SLATE)

    # Teal separator line
    p2 = hdr.add_paragraph()
    _spacing(p2, before=0, after=0)
    pPr = p2._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), TEAL_HEX)
    pBdr.append(bot)
    pPr.append(pBdr)
    p2.add_run(" ").font.size = Pt(1)


def _build_footer(section):
    """Footer: company name left, page number right."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.clear()
    _spacing(p, before=4, after=0)

    # Teal top separator
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top  = OxmlElement("w:top")
    top.set(qn("w:val"),   "single")
    top.set(qn("w:sz"),    "6")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), TEAL_HEX)
    pBdr.append(top)
    pPr.append(pBdr)

    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), "8640")
    tabs.append(tab)
    pPr.append(tabs)

    r_left = p.add_run("CloudNavision Private Limited  |  cloudnavision.com")
    _font(r_left, 8, italic=True, color=SLATE)
    r_tab = p.add_run("\t")
    _font(r_tab, 8)
    r_pg = p.add_run("Page ")
    _font(r_pg, 8, color=SLATE)
    r_num = p.add_run("")
    _font(r_num, 8, color=TEAL)
    _add_page_number(r_num)


# ── Builder ───────────────────────────────────────────────────────────────────

def build(force: bool = False):
    if not force and TEMPLATE_PATH.exists():
        try:
            if _VERSION_PATH.read_text().strip() == _TEMPLATE_VERSION:
                return
        except Exception:
            pass

    TEMPLATE_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.top_margin    = Cm(2.8)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)
    _build_header(sec)
    _build_footer(sec)

    # ── COVER TITLE BLOCK ─────────────────────────────────────────────────────
    # Navy bar
    p_title = doc.add_paragraph()
    _shade_para(p_title, NAVY_HEX)
    _spacing(p_title, before=0, after=0)
    r = p_title.add_run("  MEETING MINUTES")
    _font(r, 24, bold=True, color=WHITE)

    # Teal accent line
    p_accent = doc.add_paragraph()
    _shade_para(p_accent, TEAL_HEX)
    _spacing(p_accent, before=0, after=0)
    p_accent.add_run("").font.size = Pt(3)

    # Subtitle row
    p_sub = doc.add_paragraph()
    _shade_para(p_sub, TEAL_LT_HEX)
    _spacing(p_sub, before=0, after=0)
    r = p_sub.add_run("  {{ doc_title }}")
    _font(r, 13, bold=True, color=NAVY)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ── MEETING DETAILS TABLE ─────────────────────────────────────────────────
    _section_bar(doc, "MEETING DETAILS")

    det = doc.add_table(rows=5, cols=4)
    det.alignment = WD_TABLE_ALIGNMENT.LEFT
    _tbl_borders(det)

    col_widths = [Cm(4.0), Cm(6.5), Cm(3.5), Cm(5.0)]
    for row in det.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]

    details_data = [
        ("Meeting Title",     "{{ doc_title }}",      "Reference No.",  "MM-{{ generated_date | replace(' ', '') }}"),
        ("Date & Time",       "{{ meeting_date }}",   "Duration",       ""),
        ("Location / Platform","{{ location }}",      "Prepared By",    "CloudNavision"),
        ("Facilitated By",    "{{ facilitator }}",    "Department",     ""),
        ("Purpose",           "Process Walkthrough & Knowledge Capture", "Status", "Final"),
    ]
    for i, (l1, v1, l2, v2) in enumerate(details_data):
        row = det.rows[i]
        bg = TEAL_LT_HEX if i % 2 == 0 else WHITE_HEX
        _cell_bg(row.cells[0], TEAL_MED_HEX)
        _cell_bg(row.cells[1], bg)
        _cell_bg(row.cells[2], TEAL_MED_HEX)
        _cell_bg(row.cells[3], bg)
        r1 = row.cells[0].paragraphs[0].add_run(l1)
        _font(r1, 9, bold=True, color=NAVY)
        r2 = row.cells[1].paragraphs[0].add_run(v1)
        _font(r2, 9, color=DARK)
        r3 = row.cells[2].paragraphs[0].add_run(l2)
        _font(r3, 9, bold=True, color=NAVY)
        r4 = row.cells[3].paragraphs[0].add_run(v2)
        _font(r4, 9, color=DARK)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── ATTENDEES TABLE ───────────────────────────────────────────────────────
    _section_bar(doc, "ATTENDEES")

    att_tbl = doc.add_table(rows=2, cols=4)
    att_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _tbl_borders(att_tbl)

    att_col_widths = [Cm(5.5), Cm(4.5), Cm(4.5), Cm(3.5)]
    headers = ["Name", "Role / Designation", "Organisation", "Present (✓/✗)"]
    for i, hdr_txt in enumerate(headers):
        cell = att_tbl.rows[0].cells[i]
        cell.width = att_col_widths[i]
        _cell_bg(cell, NAVY_HEX)
        r = cell.paragraphs[0].add_run(hdr_txt)
        _font(r, 9, bold=True, color=WHITE)

    # Data row — attendees is rendered as a single block
    _cell_bg(att_tbl.rows[1].cells[0], TEAL_LT_HEX)
    r_att = att_tbl.rows[1].cells[0].paragraphs[0].add_run("{{ attendees }}")
    _font(r_att, 9, color=DARK)
    for ci in [1, 2, 3]:
        _cell_bg(att_tbl.rows[1].cells[ci], TEAL_LT_HEX)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── AGENDA & DISCUSSION ───────────────────────────────────────────────────
    _section_bar(doc, "AGENDA & DISCUSSION ITEMS")

    _ctrl(doc, "{%- for item in agenda_items %}")

    # Item number + title heading
    p_item = doc.add_paragraph()
    _shade_para(p_item, TEAL_LT_HEX)
    _spacing(p_item, before=10, after=0)
    r_num  = p_item.add_run("  {{ item.sequence }}.")
    _font(r_num, 11, bold=True, color=TEAL)
    r_title = p_item.add_run("  {{ item.title }}")
    _font(r_title, 11, bold=True, color=NAVY)

    # Teal bottom border below heading
    _border_bottom(p_item, TEAL_HEX, size=6)

    # Discussion notes
    p_lbl = doc.add_paragraph()
    _spacing(p_lbl, before=6, after=0)
    r = p_lbl.add_run("Discussion Notes:")
    _font(r, 9, bold=True, color=TEAL)

    p_disc = doc.add_paragraph()
    r = p_disc.add_run("{{ item.discussion_notes }}")
    _font(r, 10, color=DARK)
    _spacing(p_disc, before=2, after=4)

    # Screenshot
    _ctrl(doc, "{%- if item.screenshot %}")
    p_ss = doc.add_paragraph()
    r = p_ss.add_run("{{ item.screenshot }}")
    _font(r, 10)
    _spacing(p_ss, before=6, after=2)
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_cap.add_run("Screenshot — Item {{ item.sequence }}")
    _font(r, 9, italic=True, color=SLATE)
    _spacing(p_cap, before=0, after=6)
    _ctrl(doc, "{%- endif %}")

    # Action items table per agenda item
    _ctrl(doc, "{%- if item.action_items %}")

    p_al = doc.add_paragraph()
    _spacing(p_al, before=4, after=0)
    r = p_al.add_run("  Action Items:")
    _font(r, 9, bold=True, color=WHITE)
    _shade_para(p_al, TEAL_HEX)

    _ctrl(doc, "{%- for action in item.action_items %}")
    p_ai = doc.add_paragraph()
    _spacing(p_ai, before=1, after=1)
    r_bullet = p_ai.add_run("  ▸  ")
    _font(r_bullet, 10, color=TEAL)
    r_action = p_ai.add_run("{{ action }}")
    _font(r_action, 10, color=DARK)
    _ctrl(doc, "{%- endfor %}")

    _ctrl(doc, "{%- endif %}")

    # Separator between items
    p_sep = doc.add_paragraph()
    _spacing(p_sep, before=6, after=0)
    p_sep.add_run("").font.size = Pt(1)

    _ctrl(doc, "{%- endfor %}")

    # ── ACTION ITEMS SUMMARY ──────────────────────────────────────────────────
    doc.add_page_break()
    _section_bar(doc, "ACTION ITEMS SUMMARY")

    ai_tbl = doc.add_table(rows=1, cols=4)
    ai_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _tbl_borders(ai_tbl)

    ai_col_widths = [Cm(1.8), Cm(8.2), Cm(4.0), Cm(4.0)]
    ai_headers = ["Ref #", "Action Item", "Owner", "Status"]
    for i, hdr_txt in enumerate(ai_headers):
        cell = ai_tbl.rows[0].cells[i]
        cell.width = ai_col_widths[i]
        _cell_bg(cell, NAVY_HEX)
        r = cell.paragraphs[0].add_run(hdr_txt)
        _font(r, 9, bold=True, color=WHITE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Jinja2 loop to populate summary rows
    _ctrl(doc, "{%- for item in agenda_items %}")
    _ctrl(doc, "{%- for action in item.action_items %}")

    ai_row_tbl = doc.add_table(rows=1, cols=4)
    ai_row_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _tbl_borders(ai_row_tbl)
    for i, w in enumerate(ai_col_widths):
        ai_row_tbl.rows[0].cells[i].width = w
    _cell_bg(ai_row_tbl.rows[0].cells[0], TEAL_LT_HEX)
    _cell_bg(ai_row_tbl.rows[0].cells[1], WHITE_HEX)
    _cell_bg(ai_row_tbl.rows[0].cells[2], TEAL_LT_HEX)
    _cell_bg(ai_row_tbl.rows[0].cells[3], WHITE_HEX)

    r_ref = ai_row_tbl.rows[0].cells[0].paragraphs[0].add_run("{{ item.sequence }}.{{ loop.index }}")
    _font(r_ref, 9, bold=True, color=TEAL)
    ai_row_tbl.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_act = ai_row_tbl.rows[0].cells[1].paragraphs[0].add_run("{{ action }}")
    _font(r_act, 9, color=DARK)
    r_own = ai_row_tbl.rows[0].cells[2].paragraphs[0].add_run("")
    _font(r_own, 9, color=DARK)
    r_sta = ai_row_tbl.rows[0].cells[3].paragraphs[0].add_run("Open")
    _font(r_sta, 9, bold=True, color=ORANGE)

    _ctrl(doc, "{%- endfor %}")
    _ctrl(doc, "{%- endfor %}")

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ── NEXT STEPS & FOLLOW-UP ────────────────────────────────────────────────
    _ctrl(doc, "{%- if follow_up_sections %}")
    _section_bar(doc, "NEXT STEPS & FOLLOW-UP")
    _ctrl(doc, "{%- for section in follow_up_sections %}")
    p_st = doc.add_paragraph()
    r = p_st.add_run("{{ section.section_title }}")
    _font(r, 10, bold=True, color=NAVY)
    _spacing(p_st, before=8, after=2)
    p_sc = doc.add_paragraph()
    r = p_sc.add_run("{{ section.content_text }}")
    _font(r, 10, color=DARK)
    _ctrl(doc, "{%- endfor %}")
    _ctrl(doc, "{%- endif %}")

    # ── NEXT MEETING ──────────────────────────────────────────────────────────
    _section_bar(doc, "NEXT MEETING")

    nm_tbl = doc.add_table(rows=3, cols=2)
    nm_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _tbl_borders(nm_tbl)

    nm_col_widths = [Cm(4.5), Cm(13.5)]
    nm_rows = [
        ("Proposed Date & Time", ""),
        ("Location / Platform",  ""),
        ("Key Agenda Items",     ""),
    ]
    for i, (lbl, val) in enumerate(nm_rows):
        row = nm_tbl.rows[i]
        bg  = TEAL_LT_HEX if i % 2 == 0 else WHITE_HEX
        _cell_bg(row.cells[0], TEAL_MED_HEX)
        _cell_bg(row.cells[1], bg)
        row.cells[0].width = nm_col_widths[0]
        row.cells[1].width = nm_col_widths[1]
        r1 = row.cells[0].paragraphs[0].add_run(lbl)
        _font(r1, 9, bold=True, color=NAVY)
        r2 = row.cells[1].paragraphs[0].add_run(val)
        _font(r2, 9, color=DARK)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ── SIGN-OFF ──────────────────────────────────────────────────────────────
    _section_bar(doc, "SIGN-OFF & DISTRIBUTION")

    so_tbl = doc.add_table(rows=3, cols=3)
    so_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _tbl_borders(so_tbl)

    so_col_widths = [Cm(5.0), Cm(5.0), Cm(8.0)]
    so_headers = ["Role", "Name", "Signature & Date"]
    for i, hdr_txt in enumerate(so_headers):
        cell = so_tbl.rows[0].cells[i]
        cell.width = so_col_widths[i]
        _cell_bg(cell, NAVY_HEX)
        r = cell.paragraphs[0].add_run(hdr_txt)
        _font(r, 9, bold=True, color=WHITE)

    so_roles = ["Prepared By", "Reviewed By"]
    for i, role in enumerate(so_roles):
        row = so_tbl.rows[i + 1]
        bg  = TEAL_LT_HEX if i % 2 == 0 else WHITE_HEX
        _cell_bg(row.cells[0], TEAL_MED_HEX)
        _cell_bg(row.cells[1], bg)
        _cell_bg(row.cells[2], bg)
        for j, w in enumerate(so_col_widths):
            row.cells[j].width = w
        r1 = row.cells[0].paragraphs[0].add_run(role)
        _font(r1, 9, bold=True, color=NAVY)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── GENERATED NOTE ────────────────────────────────────────────────────────
    p_gen = doc.add_paragraph()
    p_gen.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p_gen.add_run("Generated by CloudNavision SOP Platform  |  {{ generated_date }}")
    _font(r, 8, italic=True, color=SLATE)

    doc.save(str(TEMPLATE_PATH))
    _VERSION_PATH.write_text(_TEMPLATE_VERSION)
    print(f"Meeting Minutes template v{_TEMPLATE_VERSION} written to {TEMPLATE_PATH}")


if __name__ == "__main__":
    build(force=True)
