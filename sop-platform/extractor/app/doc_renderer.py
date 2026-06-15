"""
SOP Document Renderer
Phase 7a: docxtpl template injection + LibreOffice PDF conversion + Azure Blob upload
"""
import logging
import re
import subprocess
import tempfile
import time
from datetime import date
from pathlib import Path
from typing import Optional

import requests
from docxtpl import DocxTemplate, InlineImage, RichText
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Cm

logger = logging.getLogger(__name__)

TEMPLATE_PATH = Path("/data/templates/sop_template.docx")
EXPORTS_DIR = Path("/data/exports")

# Regex for characters that are illegal in XML 1.0 (excludes \t \n \r which are fine)
_XML_ILLEGAL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


def _sanitize_text(text: str) -> str:
    """Strip XML-illegal characters that would corrupt the DOCX output."""
    if not text:
        return text
    return _XML_ILLEGAL_RE.sub("", text)


def render_sop(
    sop_id: str,
    fmt: str,                       # 'docx' or 'pdf'
    sop_data: dict,
    azure_blob_base_url: str,
    azure_sas_token: str,
    template: str = "standard",     # 'standard' | 'meeting_minutes' | 'webinar'
) -> dict:
    """
    Render a SOP document from the Word template.

    Returns:
        {"docx_url": str, "pdf_url": str | None}
        URLs are base Azure Blob URLs without SAS (safe for DB storage).
    """
    # Select template file, builder, context function, and filename prefix
    if template == "meeting_minutes":
        from app.build_template_meeting_minutes import (
            build as _build_tpl,
            TEMPLATE_PATH as _tmpl_path,
        )
        _context_fn = _build_context_meeting_minutes
        doc_prefix  = "meeting_minutes"
        run_post_process = False
    elif template == "webinar":
        from app.build_template_webinar import (
            build as _build_tpl,
            TEMPLATE_PATH as _tmpl_path,
        )
        _context_fn = _build_context_webinar
        doc_prefix  = "webinar"
        run_post_process = False
    else:
        from app.build_template import build as _build_tpl
        _tmpl_path = TEMPLATE_PATH
        _context_fn = None          # uses _build_context with table_registry
        doc_prefix  = "sop"
        run_post_process = True

    try:
        _build_tpl(force=False)
    except Exception as exc:
        logger.warning("Template builder failed: %s", exc)

    if not _tmpl_path.exists():
        raise FileNotFoundError(f"Template not found: {_tmpl_path}")

    export_dir = EXPORTS_DIR / sop_id
    export_dir.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory(prefix=f"sop_render_{sop_id}_") as tmp_str:
        tmp_dir = Path(tmp_str)

        tpl = DocxTemplate(str(_tmpl_path))
        table_registry: dict[str, list] = {}

        if _context_fn is not None:
            context = _context_fn(tpl, sop_data, tmp_dir, azure_sas_token=azure_sas_token)
        else:
            context = _build_context(tpl, sop_data, tmp_dir, table_registry, azure_sas_token=azure_sas_token)

        tpl.render(context)

        docx_filename = f"{doc_prefix}_{sop_id}.docx"
        docx_path = export_dir / docx_filename
        tpl.save(str(docx_path))

        if run_post_process:
            if table_registry:
                _inject_tables(docx_path, table_registry)
            _inject_toc_links(docx_path)

        logger.info("Rendered DOCX: %s (%.1f KB)", docx_path, docx_path.stat().st_size / 1024)

        docx_blob_path = f"exports/{sop_id}/{docx_filename}"
        docx_base_url = f"{azure_blob_base_url}/{docx_blob_path}"
        _upload_blob(
            docx_path,
            f"{docx_base_url}?{azure_sas_token}",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        logger.info("Uploaded DOCX → %s", docx_blob_path)

        pdf_base_url: Optional[str] = None

        if fmt == "pdf":
            pdf_path = _convert_to_pdf(docx_path, export_dir)
            logger.info("PDF created: %s (%.1f KB)", pdf_path, pdf_path.stat().st_size / 1024)

            pdf_filename = pdf_path.name
            pdf_blob_path = f"exports/{sop_id}/{pdf_filename}"
            pdf_base_url = f"{azure_blob_base_url}/{pdf_blob_path}"
            _upload_blob(pdf_path, f"{pdf_base_url}?{azure_sas_token}", "application/pdf")
            logger.info("Uploaded PDF → %s", pdf_blob_path)

    return {"docx_url": docx_base_url, "pdf_url": pdf_base_url}


def _section_content(tpl: DocxTemplate, section: dict, table_registry: dict):
    """Return a RichText for a section. Table sections use a placeholder string."""
    content_type = str(section.get("content_type") or "text")
    if "." in content_type:
        content_type = content_type.split(".")[-1]

    text = _sanitize_text(section.get("content_text") or "")
    json_data = section.get("content_json")

    if content_type == "table" and isinstance(json_data, list) and json_data:
        rows_data = [r for r in json_data if isinstance(r, dict)]
        if rows_data:
            placeholder = f"__TBLPH_{len(table_registry):03d}__"
            table_registry[placeholder] = rows_data
            rt = RichText()
            rt.add(placeholder)
            return rt

    if content_type == "list":
        rt = RichText()
        items = json_data if isinstance(json_data, list) else ([json_data] if json_data else [text])
        for i, item in enumerate(items):
            item_str = _sanitize_text(str(item)) if item is not None else ""
            if i > 0:
                rt.xml += "<w:r><w:br/></w:r>"
            rt.add(f"•  {item_str}", font="Calibri", size=22)
        return rt

    rt = RichText()
    rt.add(text, font="Calibri", size=22)
    return rt


def _set_cell_shd(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _add_tbl_borders(tbl) -> None:
    tblPr = tbl._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "D1D5DB")
        borders.append(b)
    tblPr.append(borders)


def _inject_tables(docx_path: Path, table_registry: dict[str, list]) -> None:
    """Open the saved docx, find placeholder paragraphs, replace with Word tables."""
    from docx import Document as DocxDoc

    doc = DocxDoc(str(docx_path))
    body = doc.element.body

    for para in doc.paragraphs:
        text = para.text.strip()
        if text not in table_registry:
            continue
        rows_data = table_registry[text]
        headers = list(rows_data[0].keys())
        n_cols = len(headers)
        col_w = Inches(5.3 / n_cols)

        # Build the table
        tbl = doc.add_table(rows=1 + len(rows_data), cols=n_cols)
        _add_tbl_borders(tbl)

        # Header row — orange bg, white bold
        for i, h in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            cell.width = col_w
            _set_cell_shd(cell, "E85C1A")
            run = cell.paragraphs[0].add_run(h.replace("_", " ").upper())
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)
            run.font.name = "Calibri"

        # Data rows — alternating bg
        for r_idx, row_data in enumerate(rows_data):
            bg = "F8F9FA" if r_idx % 2 == 0 else "FFFFFF"
            for c_idx, h in enumerate(headers):
                cell = tbl.rows[r_idx + 1].cells[c_idx]
                cell.width = col_w
                _set_cell_shd(cell, bg)
                run = cell.paragraphs[0].add_run(str(row_data.get(h) or ""))
                run.font.size = Pt(9)
                run.font.name = "Calibri"

        # Move the new table element to replace the placeholder paragraph
        tbl_element = tbl._tbl
        body.remove(tbl_element)          # add_table appended to body; move it
        para._element.addprevious(tbl_element)
        para._element.getparent().remove(para._element)

    doc.save(str(docx_path))


# ── TOC Hyperlinks ────────────────────────────────────────────────────────────

def _add_para_bookmark(para, name: str, bm_id: int) -> None:
    """Insert a named bookmark at the start of a heading paragraph."""
    p = para._p
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), str(bm_id))
    bm_start.set(qn("w:name"), name)
    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(bm_id))

    # Insert immediately after pPr (or at position 0 if no pPr)
    children = list(p)
    pPr = p.find(qn("w:pPr"))
    insert_pos = children.index(pPr) + 1 if pPr is not None else 0
    p.insert(insert_pos, bm_start)
    p.insert(insert_pos + 1, bm_end)


def _is_toc_para(para) -> bool:
    """Return True if paragraph has the TOC right-aligned dot-leader tab at 7920 twips."""
    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        return False
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        return False
    for tab in tabs.findall(qn("w:tab")):
        if (
            tab.get(qn("w:val")) == "right"
            and tab.get(qn("w:leader")) == "dot"
            and tab.get(qn("w:pos")) == "7920"
        ):
            return True
    return False


def _add_toc_hyperlink(para, anchor: str) -> None:
    """Wrap all runs in a TOC paragraph inside a w:hyperlink element pointing to anchor."""
    p = para._p
    runs = p.findall(qn("w:r"))
    if not runs:
        return

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)

    # Insert hyperlink at position of first run, then move all runs into it
    children = list(p)
    first_pos = children.index(runs[0])
    p.insert(first_pos, hyperlink)
    for r in runs:
        p.remove(r)
        hyperlink.append(r)


def _inject_toc_links(docx_path: Path) -> None:
    """
    Post-process rendered DOCX:
    1. Add named bookmarks to all Heading 1/2/3 paragraphs.
    2. Wrap TOC entry paragraphs (dot-leader tab at 7920) in w:hyperlink elements.
    Headings and TOC entries are matched by normalising whitespace in their text.
    """
    try:
        from docx import Document as DocxDoc

        doc = DocxDoc(str(docx_path))
        all_paras = list(doc.paragraphs)
        heading_styles = {"Heading 1", "Heading 2", "Heading 3"}

        # ── Pass 1: assign bookmarks to heading paragraphs ────────────────────
        bm_id = 200  # start high to avoid conflicts with Word's auto-bookmarks
        bm_map: dict[str, str] = {}  # normalised_text → bookmark_name

        for para in all_paras:
            style = para.style.name if para.style else ""
            if style not in heading_styles:
                continue
            text = " ".join(para.text.split())
            if not text:
                continue
            bm_name = f"_soptoc{bm_id}"
            bm_map[text] = bm_name
            _add_para_bookmark(para, bm_name, bm_id)
            bm_id += 1

        if not bm_map:
            return  # no headings found — skip (e.g. empty SOP)

        # ── Pass 2: add hyperlinks to TOC paragraphs ─────────────────────────
        for para in all_paras:
            if not _is_toc_para(para):
                continue
            raw_text = para.text
            # Strip tab character and anything after it (dot leaders + empty page ref)
            title_part = raw_text.split("\t")[0]
            normalised = " ".join(title_part.split())
            if not normalised:
                continue
            bm_name = bm_map.get(normalised)
            if bm_name:
                _add_toc_hyperlink(para, bm_name)

        doc.save(str(docx_path))
        logger.info("Injected TOC hyperlinks (%d bookmarks)", len(bm_map))

    except Exception as exc:
        logger.warning("TOC link injection failed (non-fatal): %s", exc)


# ── Generic-template context builders ────────────────────────────────────────

def _build_context_meeting_minutes(
    tpl: DocxTemplate, sop_data: dict, tmp_dir: Path, azure_sas_token: str = ""
) -> dict:
    """Context for the Meeting Minutes template."""
    steps_raw = sop_data.get("steps", [])
    agenda_items = []
    for step in steps_raw:
        screenshot = None
        ann_url = step.get("annotated_screenshot_url") or step.get("screenshot_url")
        if ann_url:
            screenshot = _download_inline_image(tpl, ann_url, tmp_dir, step.get("id", "unknown"))
        agenda_items.append({
            "sequence":        step.get("sequence", ""),
            "title":           _sanitize_text(step.get("title") or ""),
            "discussion_notes": _sanitize_text(step.get("description") or ""),
            "action_items":    [_sanitize_text(str(s)) for s in (step.get("sub_steps") or []) if s],
            "screenshot":      screenshot,
        })

    all_sections = sop_data.get("sections") or []
    follow_up = [
        {
            "section_title": _sanitize_text(s.get("section_title") or ""),
            "content_text":  _sanitize_text(s.get("content_text") or ""),
        }
        for s in all_sections
    ]

    today = date.today().strftime("%d %b %Y")
    return {
        "doc_title":         _sanitize_text(sop_data.get("process_name") or sop_data.get("sop_title") or ""),
        "meeting_date":      _sanitize_text(sop_data.get("meeting_date") or today),
        "location":          "",
        "facilitator":       _sanitize_text(sop_data.get("client_name") or ""),
        "generated_date":    today,
        "attendees":         "",
        "agenda_items":      agenda_items,
        "follow_up_sections": follow_up,
    }


def _build_context_webinar(
    tpl: DocxTemplate, sop_data: dict, tmp_dir: Path, azure_sas_token: str = ""
) -> dict:
    """Context for the Webinar Summary template."""
    steps_raw    = sop_data.get("steps", [])
    all_sections = sop_data.get("sections") or []

    # First pre-section (display_order < 50) → overview; rest → resources
    pre_sections  = [s for s in all_sections if (s.get("display_order") or 0) < 50]
    post_sections = [s for s in all_sections if (s.get("display_order") or 0) >= 50]
    overview_text = _sanitize_text((pre_sections[0].get("content_text") or "") if pre_sections else "")
    resource_secs = pre_sections[1:] + post_sections

    topics = []
    for step in steps_raw:
        screenshot = None
        ann_url = step.get("annotated_screenshot_url") or step.get("screenshot_url")
        if ann_url:
            screenshot = _download_inline_image(tpl, ann_url, tmp_dir, step.get("id", "unknown"))
        topics.append({
            "number":     step.get("sequence", ""),
            "title":      _sanitize_text(step.get("title") or ""),
            "content":    _sanitize_text(step.get("description") or ""),
            "key_points": [_sanitize_text(str(s)) for s in (step.get("sub_steps") or []) if s],
            "screenshot": screenshot,
        })

    today = date.today().strftime("%d %b %Y")
    return {
        "doc_title":       _sanitize_text(sop_data.get("process_name") or sop_data.get("sop_title") or ""),
        "session_date":    _sanitize_text(sop_data.get("meeting_date") or today),
        "presenter":       _sanitize_text(sop_data.get("client_name") or ""),
        "generated_date":  today,
        "overview_text":   overview_text,
        "topics":          topics,
        "resource_sections": [
            {
                "section_title": _sanitize_text(s.get("section_title") or ""),
                "content_text":  _sanitize_text(s.get("content_text") or ""),
            }
            for s in resource_secs
        ],
    }


# ── Context builder ───────────────────────────────────────────────────────────

def _build_context(tpl: DocxTemplate, sop_data: dict, tmp_dir: Path, table_registry: dict | None = None, azure_sas_token: str = "") -> dict:
    """Build the Jinja2 context dict for docxtpl."""
    steps_raw = sop_data.get("steps", [])
    steps_ctx = []

    for step in steps_raw:
        screenshot = None
        ann_url = step.get("annotated_screenshot_url") or step.get("screenshot_url")
        if ann_url:
            screenshot = _download_inline_image(tpl, ann_url, tmp_dir, step.get("id", "unknown"))

        sub_steps = [_sanitize_text(str(s)) for s in (step.get("sub_steps") or []) if s is not None]
        description = _sanitize_text(step.get("description") or "")
        seq = step.get("sequence", "")
        callouts = sorted(step.get("callouts") or [], key=lambda c: c.get("callout_number") or 0)

        # Append "(Screenshot N, callout N)" inline to the matching sub_step sentence.
        # Callout 1 → sub_step[0], callout 2 → sub_step[1], etc.
        # Extra callouts beyond the sub_step count are appended to the description.
        for idx, c in enumerate(callouts):
            num = c.get("callout_number") or (idx + 1)
            ref = f"(Screenshot {seq}, callout {num})"
            if idx < len(sub_steps):
                text = sub_steps[idx].rstrip(".")
                sub_steps[idx] = f"{text}. {ref}"
            else:
                # More callouts than sub_steps — attach to description
                description = description.rstrip(".") + f". {ref}"

        steps_ctx.append({
            "sequence": seq,
            "title": _sanitize_text(step.get("title") or ""),
            "description": description,
            "sub_steps": sub_steps,
            "screenshot": screenshot,
            "callouts": [],  # all references now inline — suppress the after-screenshot block
        })

    # Split sections: display_order < 50 before procedure, >= 50 after
    all_sections = sop_data.get("sections") or []
    raw_pre  = [s for s in all_sections if (s.get("display_order") or 0) < 50]
    raw_post = [s for s in all_sections if (s.get("display_order") or 0) >= 50]

    # First 5 pre-sections → sub-items of "1 Procedure Description"
    # Next 2 pre-sections → sections 2 and 3 (Training Prerequisites, Software Applications)
    # Remaining pre-sections + all post-sections → numbered after Detailed Procedure (6, 7, ...)
    _PROC_SUB_COUNT = 5
    _OTHER_PRE_COUNT = 2
    raw_proc_subs   = raw_pre[:_PROC_SUB_COUNT]
    raw_other_pre   = raw_pre[_PROC_SUB_COUNT:_PROC_SUB_COUNT + _OTHER_PRE_COUNT]
    raw_sections_post = raw_pre[_PROC_SUB_COUNT + _OTHER_PRE_COUNT:] + raw_post

    procedure_sub_sections = []
    for s in raw_proc_subs:
        procedure_sub_sections.append({
            "section_title": _sanitize_text(s.get("section_title") or ""),
            "content_text": _section_content(tpl, s, table_registry if table_registry is not None else {}),
        })

    sec_num = 2  # section 1 is "Procedure Description" parent
    other_pre_sections = []
    for s in raw_other_pre:
        other_pre_sections.append({
            "num": str(sec_num),
            "section_title": _sanitize_text(s.get("section_title") or ""),
            "content_text": _section_content(tpl, s, table_registry if table_registry is not None else {}),
        })
        sec_num += 1

    pm_section_num = str(sec_num); sec_num += 1  # section 4
    dp_section_num = str(sec_num); sec_num += 1  # section 5

    sections_post = []
    for s in raw_sections_post:
        sections_post.append({
            "num": str(sec_num),
            "section_title": _sanitize_text(s.get("section_title") or ""),
            "content_text": _section_content(tpl, s, table_registry if table_registry is not None else {}),
        })
        sec_num += 1

    cert_section_num = str(sec_num)

    pm_config = sop_data.get("process_map_config")
    confirmed_url = pm_config.get("confirmed_url") if pm_config else None

    if confirmed_url:
        process_map = _download_confirmed_map(tpl, confirmed_url, tmp_dir, sas_token=azure_sas_token)
        if process_map is None:
            process_map = (
                _generate_swimlane_map(tpl, pm_config, steps_raw, tmp_dir)
                if pm_config and pm_config.get("lanes") and pm_config.get("assignments")
                else _generate_process_map(tpl, steps_raw, tmp_dir)
            )
    elif pm_config and pm_config.get("lanes") and pm_config.get("assignments"):
        process_map = _generate_swimlane_map(tpl, pm_config, steps_raw, tmp_dir)
    else:
        process_map = _generate_process_map(tpl, steps_raw, tmp_dir)

    today = date.today().strftime("%d %b %Y")
    cover_page = _generate_cover_page(tpl, sop_data, tmp_dir, today)

    # ── Build TOC entries ─────────────────────────────────────────────────────
    # is_sub=False → numbered main section
    # is_sub=True  → indented sub-item under Procedure Description (no number)
    toc_entries = []
    toc_entries.append({"num": "1", "title": "Procedure Description", "is_sub": False})
    for s in procedure_sub_sections:
        toc_entries.append({"num": "", "title": s["section_title"], "is_sub": True})
    for s in other_pre_sections:
        toc_entries.append({"num": s["num"], "title": s["section_title"], "is_sub": False})
    toc_entries.append({"num": pm_section_num, "title": "Process Map", "is_sub": False})
    toc_entries.append({"num": dp_section_num, "title": "Detailed Procedure", "is_sub": False})
    for s in sections_post:
        toc_entries.append({"num": s["num"], "title": s["section_title"], "is_sub": False})
    toc_entries.append({"num": cert_section_num, "title": "SOP Author/Reviewer/Approver Certification", "is_sub": False})

    return {
        "cover_page": cover_page,
        "sop_title": _sanitize_text(sop_data.get("sop_title") or ""),
        "client_name": _sanitize_text(sop_data.get("client_name") or ""),
        "process_name": _sanitize_text(sop_data.get("process_name") or ""),
        "meeting_date": _sanitize_text(sop_data.get("meeting_date") or ""),
        "generated_date": today,
        "step_count": sop_data.get("step_count", len(steps_raw)),
        "steps": steps_ctx,
        "procedure_sub_sections": procedure_sub_sections,
        "other_pre_sections": other_pre_sections,
        "sections_post": sections_post,
        "pm_section_num": pm_section_num,
        "dp_section_num": dp_section_num,
        "cert_section_num": cert_section_num,
        "process_map": process_map,
        "toc_entries": toc_entries,
    }


_COVER_TEMPLATE = Path(__file__).parent / "assets" / "cover_template.jpg"


def _generate_cover_page(
    tpl: DocxTemplate,
    sop_data: dict,
    tmp_dir: Path,
    today: str,
) -> Optional[InlineImage]:
    """
    Use the static cover template as-is, resized to A4 (2480×3508 px at 300 DPI).
    The image fills the full page — section1 has zero margins.
    """
    try:
        from PIL import Image

        if _COVER_TEMPLATE.exists():
            img = Image.open(str(_COVER_TEMPLATE)).convert("RGB")
        else:
            logger.warning("Cover template not found at %s — using generated fallback", _COVER_TEMPLATE)
            img = _build_cover_base()

        img = img.resize((2480, 3508), Image.LANCZOS)

        cover_path = tmp_dir / "cover_page.jpg"
        img.save(str(cover_path), "JPEG", quality=95, optimize=True)
        return InlineImage(tpl, str(cover_path), width=Cm(21), height=Cm(29.7))

    except Exception as exc:
        logger.warning("Cover page generation failed: %s", exc)
        return None


def _build_cover_base() -> "Image.Image":
    """Build the Infomate cover template with Pillow (used when cover_template.jpg is absent)."""
    from PIL import Image, ImageDraw, ImageFont

    W, H = 1260, 1944
    WHITE  = (255, 255, 255)
    ORANGE = (232, 92, 26)
    SHADOW = (68, 64, 108)
    TEXT_DIM = (215, 213, 235)
    TEXT_MID = (175, 172, 205)

    split_x   = int(W * 0.38)
    panel_top = 120
    panel_bot = int(H * 0.71)

    img  = Image.new("RGB", (W, H), WHITE)
    draw = ImageDraw.Draw(img)

    for y in range(H):
        t = y / H
        draw.line([(0, y), (W, y)], fill=(int(118 - t * 32), int(120 - t * 36), int(166 - t * 26)))

    draw.rectangle([(9, panel_top + 10), (split_x + 9, panel_bot + 10)], fill=SHADOW)

    dec = 132
    draw.rounded_rectangle(
        [(split_x - dec // 2, panel_top - dec // 2), (split_x + dec // 2, panel_top + dec // 2)],
        radius=30, fill=ORANGE,
    )
    draw.rectangle([(0, panel_top), (split_x, panel_bot)], fill=WHITE)

    try:
        fnt_tiny  = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",      24)
        fnt_small = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",      28)
        fnt_logo  = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 40)
    except Exception:
        fnt_tiny = fnt_small = fnt_logo = ImageFont.load_default()

    def _bb(text, font):
        bb = draw.textbbox((0, 0), text, font=font)
        return bb[2] - bb[0], bb[3] - bb[1]

    right_pad = split_x + 40
    jk_text = "A John Keells Company"
    jk_w, _ = _bb(jk_text, fnt_small)
    draw.text((W - jk_w - 30, 22), jk_text, font=fnt_small, fill=TEXT_DIM)
    draw.line([(right_pad, 72), (W - 20, 72)], fill=TEXT_MID, width=1)
    draw.line([(right_pad, H - 78), (W - 20, H - 78)], fill=TEXT_MID, width=1)
    copy_text = "2018 Infomate Private Limited"
    copy_w, _ = _bb(copy_text, fnt_tiny)
    draw.text((W - copy_w - 30, H - 60), copy_text, font=fnt_tiny, fill=TEXT_DIM)
    draw.rectangle([(W - 85, H - 48), (W, H)], fill=WHITE)

    logo_y = panel_bot - 118
    draw.rounded_rectangle([(42, logo_y), (90, logo_y + 48)], radius=11, fill=ORANGE)
    draw.text((106, logo_y + 4), "infomate", font=fnt_logo, fill=(65, 65, 75))

    return img


def _crop_taskbar(img: "PILImage.Image") -> "PILImage.Image":
    """
    Detect and remove OS chrome (title bar at top + taskbar at bottom) from screenshots.
    Uses row-mean brightness transitions; only crops when the boundary is clearly in the
    expected OS-chrome zone (top 8 % / bottom 12 %) to avoid touching content.
    """
    try:
        import numpy as np

        w, h = img.size
        if h < 300:
            return img

        arr = np.array(img.convert("L"))   # greyscale (h, w)

        # ── Bottom taskbar ────────────────────────────────────────────────
        check_h = min(120, max(40, h * 15 // 100))
        bottom = arr[h - check_h:, :]
        row_means_b = bottom.mean(axis=1)
        diffs_b = np.abs(np.diff(row_means_b))
        threshold_b = max(8.0, float(diffs_b.std()) * 1.0)
        jumps_b = np.where(diffs_b > threshold_b)[0]
        crop_bottom = h
        if jumps_b.size > 0:
            boundary_y = h - check_h + int(jumps_b[0]) + 1
            if boundary_y >= h * 0.88:
                crop_bottom = boundary_y

        # ── Top title bar ─────────────────────────────────────────────────
        # Title bars are a uniform dark band in the top ~8 %
        check_top = min(100, max(20, h * 8 // 100))
        top = arr[:check_top, :]
        row_means_t = top.mean(axis=1)
        diffs_t = np.abs(np.diff(row_means_t))
        threshold_t = max(10.0, float(diffs_t.std()) * 1.2)
        jumps_t = np.where(diffs_t > threshold_t)[0]
        crop_top = 0
        if jumps_t.size > 0:
            # Take the LAST jump — bottom edge of the title bar
            candidate = int(jumps_t[-1]) + 1
            if candidate <= h * 0.08:
                # Only crop if that band is actually dark (mean brightness < 80)
                if row_means_t[:candidate].mean() < 80:
                    crop_top = candidate

        if crop_top > 0 or crop_bottom < h:
            return img.crop((0, crop_top, w, crop_bottom))
    except Exception:
        pass  # never break the render over a crop heuristic
    return img


def _download_inline_image(
    tpl: DocxTemplate,
    url: str,
    tmp_dir: Path,
    step_id: str,
) -> Optional[InlineImage]:
    """
    Download a screenshot, auto-crop the OS taskbar, resize to max 1400 px wide,
    and save as JPEG. Keeps DOCX size small and render time fast.
    """
    try:
        from PIL import Image as PILImage
        import io

        resp = requests.get(url, timeout=30)
        resp.raise_for_status()

        img = PILImage.open(io.BytesIO(resp.content)).convert("RGB")
        img = _crop_taskbar(img)

        max_w = 1400
        if img.width > max_w:
            ratio = max_w / img.width
            img = img.resize((max_w, int(img.height * ratio)), PILImage.LANCZOS)

        img_path = tmp_dir / f"screenshot_{step_id}.jpg"
        img.save(str(img_path), "JPEG", quality=85, optimize=True)

        return InlineImage(tpl, str(img_path), width=Inches(5.5))
    except Exception as exc:
        logger.warning("Could not download screenshot for step %s: %s", step_id, exc)
        return None


def _generate_process_map(
    tpl: DocxTemplate,
    steps: list[dict],
    tmp_dir: Path,
) -> Optional[InlineImage]:
    """
    Generate a sequential process map PNG using Pillow and return as InlineImage.
    Draws orange-accented step boxes with downward arrows between them.
    """
    if not steps:
        return None
    try:
        from PIL import Image, ImageDraw, ImageFont

        # Layout
        IMG_W     = 1400
        PADDING   = 50
        BOX_H     = 82
        BOX_R     = 12       # corner radius
        ARROW_H   = 34
        CIRCLE_R  = 24
        HEADER_H  = 72

        # Palette
        ORANGE      = (232, 92, 26)
        LIGHT_GREY  = (245, 245, 245)
        BORDER      = (204, 204, 204)
        TEXT_DARK   = (26, 26, 26)
        WHITE       = (255, 255, 255)

        n = len(steps)
        total_h = HEADER_H + PADDING + n * BOX_H + (n - 1) * ARROW_H + PADDING

        img  = Image.new("RGB", (IMG_W, total_h), WHITE)
        draw = ImageDraw.Draw(img)

        # Fonts (fall back to default pixel font if DejaVu not present)
        try:
            fnt_head = ImageFont.truetype(
                "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 24
            )
            fnt_body = ImageFont.truetype(
                "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 19
            )
            fnt_num  = ImageFont.truetype(
                "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 21
            )
        except Exception:
            fnt_head = fnt_body = fnt_num = ImageFont.load_default()

        # Header bar
        draw.rectangle([(0, 0), (IMG_W, HEADER_H)], fill=ORANGE)
        draw.text((PADDING, HEADER_H // 2 - 14), "Process Flow", font=fnt_head, fill=WHITE)

        y      = HEADER_H + PADDING
        box_x1 = PADDING
        box_x2 = IMG_W - PADDING
        mid_x  = IMG_W // 2

        for i, step in enumerate(steps):
            is_last = i == n - 1

            # Step box
            draw.rounded_rectangle(
                [(box_x1, y), (box_x2, y + BOX_H)],
                radius=BOX_R,
                fill=LIGHT_GREY,
                outline=ORANGE if is_last else BORDER,
                width=2,
            )

            # Numbered circle on the left
            cx = box_x1 + PADDING // 2 + CIRCLE_R
            cy = y + BOX_H // 2
            draw.ellipse(
                [(cx - CIRCLE_R, cy - CIRCLE_R), (cx + CIRCLE_R, cy + CIRCLE_R)],
                fill=ORANGE,
            )
            num_str = str(step.get("sequence", i + 1))
            draw.text((cx - CIRCLE_R // 2 - 1, cy - CIRCLE_R // 2 + 1), num_str, font=fnt_num, fill=WHITE)

            # Step title
            title    = step.get("title", "")
            title    = title[:72] + ("…" if len(title) > 72 else "")
            text_x   = cx + CIRCLE_R + 16
            text_y   = y + BOX_H // 2 - 12
            draw.text((text_x, text_y), title, font=fnt_body, fill=TEXT_DARK)

            y += BOX_H

            # Arrow between steps
            if not is_last:
                arrow_tip = y + ARROW_H
                draw.line([(mid_x, y), (mid_x, arrow_tip - 10)], fill=BORDER, width=2)
                draw.polygon(
                    [(mid_x - 9, arrow_tip - 10), (mid_x + 9, arrow_tip - 10), (mid_x, arrow_tip)],
                    fill=BORDER,
                )
                y += ARROW_H

        map_path = tmp_dir / "process_map.jpg"
        img.save(str(map_path), "JPEG", quality=92, optimize=True)
        return InlineImage(tpl, str(map_path), width=Inches(5.5))

    except Exception as exc:
        logger.warning("Could not generate process map: %s", exc)
        return None


def _hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))  # type: ignore[return-value]


def _generate_swimlane_map(
    tpl: DocxTemplate,
    config: dict,
    steps: list[dict],
    tmp_dir: Path,
) -> Optional[InlineImage]:
    """
    Generate a swim-lane process map PNG from process_map_config.
    Lanes are vertical columns; steps flow top-to-bottom with cross-lane arrows.
    """
    if not steps or not config:
        return None
    try:
        from PIL import Image, ImageDraw, ImageFont

        lanes = config.get("lanes", [])
        assignments = config.get("assignments", [])
        if not lanes or not assignments:
            return None

        step_by_id = {s.get("id"): s for s in steps}
        lane_idx = {l["id"]: i for i, l in enumerate(lanes)}

        LANE_W   = 300
        ROW_H    = 110
        BOX_W    = 260
        BOX_H    = 64
        HEADER_H = 58
        MARGIN   = 20
        CIRCLE_R = 18

        n_lanes = len(lanes)
        n_rows  = len(assignments)

        IMG_W = MARGIN + n_lanes * LANE_W + MARGIN
        IMG_H = MARGIN + HEADER_H + n_rows * ROW_H + MARGIN

        WHITE  = (255, 255, 255)
        LIGHT  = (248, 250, 252)
        ALT    = (241, 245, 249)
        BORDER = (203, 213, 225)
        DARK   = (15, 23, 42)
        ARROW  = (148, 163, 184)
        AMBER  = (217, 119, 6)
        CREAM  = (255, 251, 235)

        img  = Image.new("RGB", (IMG_W, IMG_H), WHITE)
        draw = ImageDraw.Draw(img)

        try:
            fnt_head = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 18)
            fnt_body = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 15)
            fnt_num  = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 14)
            fnt_small = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 13)
        except Exception:
            fnt_head = fnt_body = fnt_num = fnt_small = ImageFont.load_default()

        def text_size(text: str, font) -> tuple[int, int]:
            """Return (width, height) of rendered text using textbbox."""
            try:
                bb = draw.textbbox((0, 0), text, font=font)
                return bb[2] - bb[0], bb[3] - bb[1]
            except Exception:
                return len(text) * 8, 16

        def draw_centered(text: str, cx: int, cy: int, font, fill):
            """Draw text centered at (cx, cy)."""
            w, h = text_size(text, font)
            draw.text((cx - w // 2, cy - h // 2), text, font=font, fill=fill)

        def wrap_title(title: str, max_w: int, font) -> list[str]:
            """Wrap title to fit within max_w pixels, max 2 lines."""
            words = title.split()
            lines: list[str] = []
            current = ""
            for word in words:
                test = (current + " " + word).strip()
                w, _ = text_size(test, font)
                if w > max_w and current:
                    lines.append(current)
                    current = word
                    if len(lines) >= 2:
                        break
                else:
                    current = test
            if current and len(lines) < 2:
                lines.append(current)
            # Truncate last line if still too wide
            if lines and text_size(lines[-1], font)[0] > max_w:
                while lines[-1] and text_size(lines[-1] + "…", font)[0] > max_w:
                    lines[-1] = lines[-1][:-1]
                lines[-1] += "…"
            return lines or [""]

        # ── Lane backgrounds + headers ────────────────────────────────────────
        for i, lane in enumerate(lanes):
            lx = MARGIN + i * LANE_W
            ly = MARGIN
            bg = LIGHT if i % 2 == 0 else ALT
            draw.rectangle([(lx, ly), (lx + LANE_W, ly + HEADER_H + n_rows * ROW_H)], fill=bg)
            color_rgb = _hex_to_rgb(lane.get("color", "#3B82F6"))
            draw.rectangle([(lx, ly), (lx + LANE_W, ly + HEADER_H)], fill=color_rgb)
            name = lane.get("name", f"Lane {i + 1}")
            draw_centered(name, lx + LANE_W // 2, ly + HEADER_H // 2, fnt_head, WHITE)
            if i > 0:
                draw.line([(lx, MARGIN), (lx, IMG_H - MARGIN)], fill=BORDER, width=1)

        draw.rectangle([(MARGIN, MARGIN), (IMG_W - MARGIN, IMG_H - MARGIN)], outline=BORDER, width=2)

        def box_center(row: int, lane_id: str) -> tuple[int, int]:
            li = lane_idx.get(lane_id, 0)
            cx = MARGIN + li * LANE_W + LANE_W // 2
            cy = MARGIN + HEADER_H + row * ROW_H + ROW_H // 2
            return cx, cy

        # ── Arrows (drawn behind boxes) ───────────────────────────────────────
        for i, asgn in enumerate(assignments[:-1]):
            next_asgn = assignments[i + 1]
            x1, y1 = box_center(i, asgn["lane_id"])
            x2, y2 = box_center(i + 1, next_asgn["lane_id"])
            is_decision_from = asgn.get("is_decision", False)
            is_decision_to   = next_asgn.get("is_decision", False)
            half_from = (BOX_H // 2 + 8) if is_decision_from else BOX_H // 2
            half_to   = (BOX_H // 2 + 8) if is_decision_to   else BOX_H // 2

            ay_from = y1 + half_from
            ay_to   = y2 - half_to - 4
            mid_y   = y1 + ROW_H // 2

            if x1 == x2:
                draw.line([(x1, ay_from), (x2, ay_to)], fill=ARROW, width=2)
            else:
                draw.line([(x1, ay_from), (x1, mid_y)], fill=ARROW, width=2)
                draw.line([(x1, mid_y), (x2, mid_y)], fill=ARROW, width=2)
                draw.line([(x2, mid_y), (x2, ay_to)], fill=ARROW, width=2)

            draw.polygon([(x2 - 7, ay_to), (x2 + 7, ay_to), (x2, ay_to + 12)], fill=ARROW)

        # ── Step boxes / diamonds ─────────────────────────────────────────────
        for i, asgn in enumerate(assignments):
            step = step_by_id.get(asgn.get("step_id"), {})
            cx, cy = box_center(i, asgn["lane_id"])
            lane   = lanes[lane_idx.get(asgn["lane_id"], 0)]
            color_rgb = _hex_to_rgb(lane.get("color", "#3B82F6"))
            seq_num   = step.get("sequence", i + 1)
            raw_title = step.get("title") or ""

            if asgn.get("is_decision"):
                # Diamond — larger to fit text
                hw, hh = BOX_W // 2, BOX_H // 2 + 10
                draw.polygon([(cx, cy - hh), (cx + hw, cy), (cx, cy + hh), (cx - hw, cy)], fill=CREAM, outline=AMBER, width=2)
                # Sequence label top-left of diamond
                seq_label = f"{seq_num}."
                draw_centered(seq_label, cx - hw // 2, cy - hh // 2 + 4, fnt_num, AMBER)
                # Title wrapped inside diamond (max ~BOX_W - 60px safe inner area)
                title_lines = wrap_title(raw_title, BOX_W - 60, fnt_small)
                line_h = 16
                total_h = len(title_lines) * line_h
                start_y = cy - total_h // 2
                for li, line in enumerate(title_lines):
                    draw_centered(line, cx, start_y + li * line_h, fnt_small, DARK)
            else:
                # Rounded rectangle
                bx = cx - BOX_W // 2
                by = cy - BOX_H // 2
                draw.rounded_rectangle([(bx, by), (bx + BOX_W, by + BOX_H)], radius=8, fill=WHITE, outline=color_rgb, width=2)
                # Number circle
                ccx = bx + 12 + CIRCLE_R
                draw.ellipse([(ccx - CIRCLE_R, cy - CIRCLE_R), (ccx + CIRCLE_R, cy + CIRCLE_R)], fill=color_rgb)
                draw_centered(str(seq_num), ccx, cy, fnt_num, WHITE)
                # Title — wrap to fit in box width minus circle area
                title_area_w = BOX_W - (12 + CIRCLE_R * 2 + 12)
                title_lines = wrap_title(raw_title, title_area_w, fnt_body)
                line_h = 18
                total_h = len(title_lines) * line_h
                tx = ccx + CIRCLE_R + 10
                start_y = cy - total_h // 2
                for li, line in enumerate(title_lines):
                    draw.text((tx, start_y + li * line_h), line, font=fnt_body, fill=DARK)

        map_path = tmp_dir / "process_map_swimlane.jpg"
        img.save(str(map_path), "JPEG", quality=92, optimize=True)
        return InlineImage(tpl, str(map_path), width=Inches(6.5))

    except Exception as exc:
        logger.warning("Could not generate swimlane map: %s", exc)
        return _generate_process_map(tpl, steps, tmp_dir)


def _download_confirmed_map(
    tpl: DocxTemplate,
    url: str,
    tmp_dir: Path,
    sas_token: str = "",
) -> Optional[InlineImage]:
    """Download the user-uploaded confirmed process map PNG and embed it in the document."""
    try:
        full_url = f"{url}?{sas_token}" if sas_token and "?" not in url else url
        resp = requests.get(full_url, timeout=30)
        resp.raise_for_status()
        from PIL import Image as PILImage
        import io as _io
        pmap = PILImage.open(_io.BytesIO(resp.content)).convert("RGB")
        map_path = tmp_dir / "process_map_confirmed.jpg"
        pmap.save(str(map_path), "JPEG", quality=92, optimize=True)
        return InlineImage(tpl, str(map_path), width=Inches(6.5))
    except Exception as exc:
        logger.warning("Could not download confirmed process map from %s: %s", url, exc)
        return None


def _convert_to_pdf(docx_path: Path, output_dir: Path) -> Path:
    """Convert a .docx to .pdf using LibreOffice headless."""
    cmd = [
        "libreoffice",
        "--headless",
        "--convert-to", "pdf",
        "--outdir", str(output_dir),
        str(docx_path),
    ]
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice conversion failed: {result.stderr[-500:]}")

    pdf_path = output_dir / docx_path.with_suffix(".pdf").name
    if not pdf_path.exists():
        raise RuntimeError(f"LibreOffice ran but PDF not found at {pdf_path}")
    return pdf_path


def _upload_blob(local_path: Path, sas_url: str, content_type: str, max_retries: int = 3) -> None:
    """PUT a file to Azure Blob Storage using a SAS URL with retry on connection errors."""
    file_size = local_path.stat().st_size
    for attempt in range(max_retries):
        try:
            with open(local_path, "rb") as f:
                resp = requests.put(
                    sas_url,
                    data=f,
                    headers={
                        "x-ms-blob-type": "BlockBlob",
                        "Content-Type": content_type,
                        "Content-Length": str(file_size),
                    },
                    timeout=300,
                )
            resp.raise_for_status()
            return
        except (requests.exceptions.ConnectionError, requests.exceptions.Timeout) as e:
            if attempt < max_retries - 1:
                logger.warning("Blob upload failed (attempt %d/%d): %s — retrying", attempt + 1, max_retries, e)
                time.sleep(5 * (attempt + 1))
            else:
                raise
